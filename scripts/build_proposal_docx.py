import os
import sys
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_docx(input_dir, output_path):
    print(f"Generating DOCX from markdown files in {input_dir}...")
    doc = Document()
    
    # 1. Set Page Margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.4)
        
    # 2. Set Default Font (SimSun / 宋体, Xiaosi / 12pt)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    # Configure Asian Font to SimSun
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # Set exact line spacing to 22 pt
    style.paragraph_format.line_spacing = Pt(22)
    style.paragraph_format.line_spacing_rule = 4 # WD_LINE_SPACING.EXACTLY
    
    # 3. Read markdown files and append to doc
    # Look for files like section_1.md, section_2.md, etc.
    files = sorted([f for f in os.listdir(input_dir) if f.endswith('.md') and not f.startswith('reference_basis')])
    
    if not files:
        print(f"Error: No section markdown files found in {input_dir}")
        sys.exit(1)
        
    for md_file in files:
        file_path = os.path.join(input_dir, md_file)
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
            # Very basic markdown to paragraph conversion
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                p = doc.add_paragraph()
                
                # Basic styling for headers (lines starting with #)
                if line.startswith('#'):
                    # Strip hashes
                    clean_line = line.lstrip('#').strip()
                    run = p.add_run(clean_line)
                    run.bold = True
                    # Headers might want to be slightly larger, but strict requirements dictate 12pt globally unless specified
                    run.font.size = Pt(12) 
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                else:
                    run = p.add_run(line)
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    
    doc.save(output_path)
    print(f"Successfully generated formatted DOCX at: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python build_proposal_docx.py <input_dir_containing_mds> <output_docx_path>")
        sys.exit(1)
        
    input_directory = sys.argv[1]
    output_docx = sys.argv[2]
    create_docx(input_directory, output_docx)
